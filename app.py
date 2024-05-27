from flask import Flask, render_template, request, jsonify, send_file, url_for
from docx import Document
import os
import requests
from dotenv import load_dotenv

load_dotenv()

app = Flask(__name__)

BASE_API_URL = os.getenv("BASE_API_URL")
FLOW_ID = os.getenv("FLOW_ID")
TWEAKS = {
    "SystemMessagePromptTemplate-H94x4": {},
    "HumanMessagePromptTemplate-XyOK9": {},
    "ChatMessagePromptTemplate-DvdJO": {},
    "ChatOpenAI-XyaTd": {},
    "ChatPromptTemplate-PODTx": {},
    "LLMChain-JjiBl": {},
    "ChatPromptTemplate-wMeK9": {},
    "LLMChain-0k057": {},
    "ChatPromptTemplate-Z4ko7": {},
    "LLMChain-8rji7": {}
}

UPLOAD_FOLDER = 'uploads'
if not os.path.exists(UPLOAD_FOLDER):
    os.makedirs(UPLOAD_FOLDER)

def run_flow(inputs: dict, flow_id: str, tweaks: dict = None) -> dict:
    api_url = f"{BASE_API_URL}/{flow_id}"
    payload = {"inputs": inputs, "tweaks": tweaks}
    response = requests.post(api_url, json=payload)
    if response.status_code != 200:
        app.logger.error(f"Error from API: {response.text}")
    return response.json()

@app.route('/')
def index():
    return render_template('index.html')

@app.route('/chat_api', methods=['POST'])
def chat_api():
    data = request.json
    inputs = {
        "detailed_description": data.get('detailed_description', ''),
        "project_name": data.get('project_name', ''),
        "project_type": data.get('project_type', ''),
        "specific_requirements": data.get('specific_requirements', ''),
        "stakeholders": data.get('stakeholders', ''),
        "timeline": data.get('timeline', '')
    }

    try:
        response = run_flow(inputs, flow_id=FLOW_ID, tweaks=TWEAKS)
        bot_response = response.get('result', {}).get('text', 'No response from bot.')
        app.logger.info(f"Bot response: {bot_response}")

        doc = Document()
        doc.add_heading('Project Management Document', 0)
        doc.add_paragraph(bot_response)

        file_path = os.path.join(UPLOAD_FOLDER, 'document.docx')
        doc.save(file_path)
        app.logger.info(f"Document saved at: {file_path}")

        if not os.path.exists(file_path):
            raise FileNotFoundError(f"File not found: {file_path}")

        return jsonify({'response': bot_response, 'file_url': url_for('uploaded_file', filename='document.docx', _external=True)})
    except Exception as e:
        app.logger.error(f"Error during document generation: {str(e)}")
        return jsonify({'error': str(e)}), 500

@app.route('/uploads/<filename>')
def uploaded_file(filename):
    file_path = os.path.join(UPLOAD_FOLDER, filename)
    if not os.path.exists(file_path):
        app.logger.error(f"Requested file does not exist: {file_path}")
        return jsonify({'error': 'File not found'}), 404
    return send_file(file_path)

if __name__ == '__main__':
    app.run(debug=True, port=5001)
